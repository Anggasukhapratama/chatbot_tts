# my_flask_app/sebayu_app/utils.py
import os
import subprocess
import tempfile
import shutil
import re
import requests
import json
import threading
import uuid
import time
from io import BytesIO
from urllib.parse import urlparse
from pathlib import Path
from typing import Optional, Tuple, List, Dict
from datetime import datetime

from faster_whisper import WhisperModel

from .config import (
    UPLOAD_DIR, ALLOWED_AUDIO, HAVE_DOCX, log, PROGRESS, DEFAULT_META, PROJECT_ROOT
)
from .database import get_db, now_str, current_program, get_today_schedule_text
from .textclean import clean_text_id  # <--- Cleaner terintegrasi

# (opsional) ambil preferensi device/compute dari env via config; fallback aman
try:
    from .config import WHISPER_DEVICE, WHISPER_COMPUTE
except Exception:
    WHISPER_DEVICE, WHISPER_COMPUTE = "auto", "float16"

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_AUDIO

# --- FFmpeg helpers ---
def ffprobe_duration(path: Path) -> float:
    try:
        cmd = [
            "ffprobe", "-v", "error", "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1", str(path)
        ]
        out = subprocess.check_output(cmd, stderr=subprocess.STDOUT).decode().strip()
        return float(out)
    except subprocess.CalledProcessError as e:
        log.error(f"ffprobe error: {e.output.decode(errors='ignore')}")
        return 0.0
    except Exception as e:
        log.error(f"ffprobe parse error: {e}")
        return 0.0

def ffmpeg_preprocess(in_path: Path) -> Path:
    out_path = in_path.with_suffix("")
    out_path = Path(str(out_path) + "__16k.wav")
    cmd = [
        "ffmpeg", "-y", "-i", str(in_path),
        "-ac", "1", "-ar", "16000", "-vn",
        "-af", "loudnorm=I=-16:TP=-2:LRA=11",
        str(out_path)
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return out_path

def ffmpeg_segment(in_path: Path, segment_seconds: int = 600) -> List[Path]:
    tempdir = Path(tempfile.mkdtemp(prefix="segments_"))
    pattern = tempdir / "part_%03d.wav"
    pre = ffmpeg_preprocess(in_path)
    cmd = [
        "ffmpeg", "-y", "-i", str(pre),
        "-f", "segment", "-segment_time", str(segment_seconds),
        "-c", "copy", str(pattern)
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    parts = sorted(tempdir.glob("part_*.wav"))
    return parts

# --- Whisper/Faster-Whisper (AUTO GPU → CPU fallback) ---
_MODEL_CACHE: dict = {}

def _get_model_cached(model_size: str, device: str, compute_type: str):
    key = (model_size, device, compute_type)
    if key not in _MODEL_CACHE:
        _MODEL_CACHE[key] = WhisperModel(model_size, device=device, compute_type=compute_type)
    return _MODEL_CACHE[key]

def choose_model(duration_sec: float, mode: str, manual_choice: str) -> str:
    if mode == "manual" and manual_choice in {"tiny", "base", "small", "medium"}:
        return manual_choice
    if duration_sec > 1800:  # 30 minutes
        return "medium"
    return "small"

def run_faster_whisper(audio_file: Path, model_size: str) -> str:
    """
    Auto-detect GPU:
      - Coba CUDA + compute_type (default: float16 / via env)
      - Kalau gagal → fallback ke CPU int8
    Bisa dipaksa via env (opsional, jika didefinisikan di config.py):
      WHISPER_DEVICE=auto|cuda|cpu
      WHISPER_COMPUTE=float16|int8_float16|int8
    """
    user_device = (WHISPER_DEVICE or "auto").lower()
    user_compute = (WHISPER_COMPUTE or "float16").lower()

    device = "cuda" if user_device not in {"cpu", "cuda"} else user_device
    compute_type = user_compute if user_compute in {"float16", "int8_float16", "int8"} else "float16"

    model = None
    if device == "cuda":
        try:
            model = _get_model_cached(model_size, "cuda", compute_type)
            log.info(f"WhisperModel loaded on CUDA ({compute_type})")
        except Exception as e:
            log.warning(f"GPU unavailable ({e}); falling back to CPU int8")
            model = _get_model_cached(model_size, "cpu", "int8")
    else:
        model = _get_model_cached(model_size, "cpu", "int8")
        log.info("WhisperModel loaded on CPU (int8)")

    domain_prompt = (
        "Sebayu FM, Diskominfo, Tegal, Slawi, Brebes, "
        "Berita Pagi, Musik Santai, Relaks Malam, Sabtu Ceria, Pemkab, notulensi rapat, agenda, keputusan."
    )
    segments, _ = model.transcribe(
        str(audio_file),
        language="id",
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
        beam_size=5,           # untuk speed bisa turunkan ke 1–3
        best_of=5,             # untuk speed bisa turunkan ke 1–2
        condition_on_previous_text=False,
        initial_prompt=domain_prompt,
    )
    texts = [seg.text.strip() for seg in segments if getattr(seg, "text", None)]
    return " ".join(texts).strip()

def transcribe_audio_pipeline(
    audio_path: Path,
    mode: str = "auto",
    manual_choice: str = "small",
    do_chunk: bool = False,
    progress=None
) -> str:
    duration = ffprobe_duration(audio_path)
    if progress: progress(20, f"Durasi terdeteksi ~{duration/60:.1f} menit")
    model_size = choose_model(duration, mode, manual_choice)
    if progress: progress(25, f"Pilih model: {model_size}")

    if do_chunk:
        if progress: progress(28, "Segmentasi audio (tiap 10 menit)")
        parts = ffmpeg_segment(audio_path, segment_seconds=600)
        n = max(1, len(parts))
        chunks_text = []
        try:
            for i, p in enumerate(parts, 1):
                share_start = 30 + int(55*(i-1)/n)
                share_end   = 30 + int(55*i/n)
                if progress: progress(share_start, f"Transkrip bagian {i}/{n}…")
                t = run_faster_whisper(p, model_size)
                chunks_text.append(f"[Bagian {i}] {t}")
                if progress: progress(share_end, f"Selesai bagian {i}/{n}")
            full_text = "\n".join(chunks_text).strip()
        finally:
            if parts:
                shutil.rmtree(parts[0].parent, ignore_errors=True)
        if progress: progress(88, "Menggabungkan teks")
        return full_text
    else:
        if progress: progress(28, "Preprocess audio")
        pre = ffmpeg_preprocess(audio_path)
        if progress: progress(35, "Transkripsi (tanpa potong)…")
        text = run_faster_whisper(pre, model_size)
        if progress: progress(88, "Finalisasi teks")
        return text

# --- Summarization lokal (tanpa OpenAI) ---
def local_summarize_bullets(text: str, max_sentences: int = 18) -> str:
    sents = re.split(r'(?<=[.!?])\s+|\n{2,}', text)
    keywords = [
        "keputusan","tindak lanjut","action","deadline","anggaran","solusi","usulan",
        "disepakati","menyetujui","ditetapkan","menetapkan","menugaskan","PIC","owner",
        "target","paling lambat","KUA","PPAS","TAPD","RKPD"
    ]
    scored = []
    for s in sents:
        s2 = s.strip()
        if not s2: continue
        score = len(s2)
        low = s2.lower()
        score += sum(28 for k in keywords if k in low)
        scored.append((score, s2))
    scored.sort(reverse=True, key=lambda x: x[0])
    bullets = [f"- {s}" for _, s in scored[:max_sentences]]
    if not bullets:
        return "- (tidak cukup konten untuk diringkas)"
    return "\n".join(bullets)

# ====== Klasifikasi Lokal Notulen (tanpa OpenAI) ======
RE_DATE = re.compile(
    r"\b(\d{1,2}[/\-.]\d{1,2}([/\-.]\d{2,4})?|"
    r"\d{1,2}\s+(jan|feb|mar|apr|mei|jun|jul|agu|sep|okt|nov|des)[a-z]*\s+\d{2,4}|"
    r"paling lambat|selambat-lambatnya|sebelum tanggal|minggu ke-\d+|akhir bulan|awal bulan)\b",
    re.I
)
RE_OWNER = re.compile(
    r"\b(PIC|owner|penanggung jawab|ditugaskan kepada|menugaskan|oleh)\b[: ]+([A-Z][a-zA-Z_. ]+)",
    re.I
)
RE_DECISION_VERB = re.compile(
    r"\b(diputuskan|memutuskan|menetapkan|ditetapkan|disepakati|menyepakati|"
    r"menyetujui|disetujui|menolak|dengan catatan)\b", re.I)
RE_ACTION_VERB   = re.compile(
    r"\b(dilakukan|ditindaklanjuti|dikerjakan|menindaklanjuti|menyelesaikan|"
    r"melakukan|follow[- ]?up|menugaskan|ditugaskan|koordinasi|menyusun|mengirim|"
    r"mengajukan|merevisi|memperbaiki|menyampaikan)\b", re.I)

BASE_KEYSETS = {
    # ================= KEPUTUSAN =================
    "keputusan": [
        # istilah eksplisit
        r"\bkeputusan\b", r"\bdiputuskan\b", r"\bmemutuskan\b", r"\bmenetapkan\b", r"\bditetapkan\b",
        r"\bdisepakati\b|\bkesepakatan\b", r"\bmenyetujui\b|\bdisetujui\b", r"\bmenolak\b", r"\bdengan catatan\b",

        # frasa rapat/formal
        r"\bhasil rapat\b", r"\bputusan rapat\b", r"\bkonklusi\b", r"\bkonklusinya\b", r"\bketetapan\b",
        r"\bpengesahan\b", r"\bdiputus\b|\bmenyepakati\b|\bmufakat\b|\bmusyawarah\b",

        # dokumen resmi
        r"\bPeraturan\b|\bPerda\b|\bPerwal\b|\bPerbup\b|\bSK\b|\bSurat Keputusan\b|\bKeppres\b|\bPerpres\b",
        r"\bditerbitkan\b|\bditetapkan melalui\b|\bdikuatkan\b|\bdisahkan\b|\bsah\b",

        # sikap/hasil musyawarah
        r"\bdipersetujui\b", r"\bdibolehkan\b", r"\bdilarang\b", r"\bditolak\b", r"\bditunda\b", r"\bdipending\b",
        r"\bdisetujui bersama\b", r"\bdengan syarat\b", r"\bdengan ketentuan\b", r"\bbersyarat\b",

        # alokasi & keputusan anggaran
        r"\bmenyetujui anggaran\b", r"\bpenetapan pagu\b", r"\bpenetapan KUA\b", r"\bpengesahan APBD\b",
        r"\bpengesahan KUA\b", r"\bpengesahan PPAS\b", r"\bpengesahan RAPBD\b", r"\bpengesahan RKPD\b",

        # keputusan personel & organisasi
        r"\bmutasi\b|\bpromosi\b|\bpemberhentian\b|\bpenunjukan\b|\bdiangkat\b|\bpengangkatan\b|\bpergantian\b",
        r"\bkeputusan pimpinan\b|\bkeputusan ketua\b|\bkeputusan dewan\b|\bhasil pleno\b",

        # lain-lain
        r"\bdinyatakan\b|\bditegaskan\b|\bdipastikan\b|\bdipilih\b|\bmemilih\b",
        r"\bdisetujui rapat\b|\bhasil sidang\b|\bketok palu\b"
    ],

    # ================= TINDAK LANJUT =================
    "tindak_lanjut": [
        # umum: aksi, follow up, target & tenggat
        r"\btindak lanjut\b", r"\bditindaklanjuti\b", r"\bmenindaklanjuti\b", r"\baksi\b|\baction\b",
        r"\bfollow[- ]?up\b", r"\bkelanjutan\b", r"\bprogress\b|\bprogres\b",
        r"\bdeadline\b|\bdue\b|\bjatuh tempo\b|\bbatas waktu\b|\bpaling lambat\b|\bselambat\b|\btarget\b|\bSLA\b|\bKPI\b",

        # penugasan & peran
        r"\bPIC\b|\bowner\b|\bpenanggung jawab\b|\bpenanggungjawab\b",
        r"\bditugaskan\b|\bmenugaskan\b|\bmenunjuk\b|\bdidelegasikan\b|\bdelegasi\b",
        r"\btanggung jawab\b|\bkoordinator\b|\bpenanggung jawab kegiatan\b",

        # koordinasi & komunikasi
        r"\bkoordinasi\b|\bberkoordinasi\b|\bkoor(d)?\b",
        r"\bkomunikasi\b|\bmenghubungi\b|\bkontak\b|\bkonfirmasi\b|\bmengonfirmasi\b",
        r"\breminder\b|\bdiingatkan\b|\bperingatan\b|\bnotifikasi\b|\bfollow[- ]?up via\b",

        # persuratan & administrasi
        r"\bmenyusun surat\b|\bmembuat surat\b|\bkonsep surat\b|\bsurat tugas\b|\bsurat undangan\b|\bnota dinas\b|\bNODIN\b",
        r"\bdisposisi\b|\bparaf\b|\bpara(f)? berjenjang\b|\btembusan\b",
        r"\bmengirim surat\b|\bmengirimkan surat\b|\bmenyampaikan surat\b|\bunggah surat\b",
        r"\bBAST\b|\bberita acara\b|\bSPT\b|\bSK\b|\bMoU\b|\bPKS\b",

        # rapat & sosialisasi
        r"\bmenjadwalkan rapat\b|\bdijadwalkan rapat\b|\bpenjadwalan\b|\batur jadwal\b",
        r"\bRapat lanjutan\b|\bFGD\b|\bbriefing\b|\bkoordinasi lintas OPD\b",
        r"\bsosialisasi\b|\bdiseminasi\b|\bpembinaan\b|\bpenguatan\b|\bpendampingan\b",

        # dokumen kerja
        r"\bmenyusun\b|\bmenyempurnakan\b|\bmelengkapi\b|\bmemperbarui\b|\bmemutakhirkan\b",
        r"\bmerevisi\b|\brevisi\b|\bperbaikan\b|\bperubahan\b|\bupdate dokumen\b",
        r"\bfinalisasi\b|\bfinal\b|\bpengesahan internal\b",
        r"\bmengarsipkan\b|\bdiarsipkan\b|\bunggah\b|\bmengunggah\b|\bupload\b|\bdiunggah\b",
        r"\bverifikasi\b|\bdipverifikasi\b|\bvalidasi\b|\bdipvalidasi\b|\bcek kelengkapan\b|\bcek dokumen\b",

        # pelaporan
        r"\bmenyusun laporan\b|\bmenyampaikan laporan\b|\blaporan kemajuan\b|\blaporan realisasi\b|\bLPJ\b|\bSPJ\b",
        r"\bmonitoring\b|\bmonitor\b|\bMonev\b|\bevaluasi\b|\breview\b|\bpenilaian\b",

        # anggaran & perencanaan
        r"\bRKPD\b|\bKUA\b|\bPPAS\b|\bRKA\b|\bDPA\b|\bPOK\b",
        r"\binput (ke )?SIPD\b|\bSIPD\b|\bSIRUP\b|\bSiRUP\b",
        r"\bpenyelarasan\b|\bsinkronisasi\b|\bpenyesuaian pagu\b|\bpenajaman program\b",

        # pengadaan
        r"\bpengadaan\b|\bPBJ\b|\bULP\b|\bLPSE\b|\be-?catalog\b|\be-katalog\b",
        r"\bSPK\b|\bkontrak\b|\bBAST\b|\bBAHN\b|\bvendor\b|\bpenyedia\b|\btagihan\b|\binvoice\b|\bpembayaran\b",

        # teknis IT / sistem
        r"\btiket\b|\bhelpdesk\b|\bissue tracking\b|\bperbaiki bug\b|\bbugfix\b|\bdeploy\b|\brelease\b|\brollout\b",
        r"\bkonfigurasi\b|\bkonfigur(asi)?\b|\bsetup\b|\binstalasi\b|\bimplementasi\b|\bgo[- ]live\b",
        r"\bdokumentasi teknis\b|\breadme\b|\bSOP\b|\bpanduan\b",

        # kegiatan lapangan
        r"\bsurvei\b|\bsurvey\b|\bverifikasi lapangan\b|\bpengukuran\b|\bpemetaan\b",
        r"\bkunjungan\b|\bvisit\b|\bpeninjauan\b|\binspeksi\b|\bcek lokasi\b",

        # legal & regulasi
        r"\bpenyusunan rancangan\b|\brancangan peraturan\b|\bperwal\b|\bperbup\b|\bperda\b",
        r"\bkonsultasi hukum\b|\bklarifikasi regulasi\b",

        # publikasi
        r"\bmenyiapkan materi\b|\bbahan paparan\b|\bslide\b|\bdeck\b|\bpress release\b|\bsiaran pers\b",
        r"\bpublikasi\b|\bunggah ke website\b|\bmedia sosial\b|\bkonten informasi\b",

        # komitmen waktu
        r"\bH\+?\d+\b|\bMinggu depan\b|\bpekan depan\b|\bbulan depan\b|\bakhir (pekan|bulan)\b|\bawal (pekan|bulan)\b",

        # frasa tugas eksplisit
        r"\bdiminta untuk\b|\bdiharapkan untuk\b|\bagar\b|\bharap\b|\bsegera\b|\bsesegera mungkin\b|\bASAP\b",
        r"\bdiperintahkan\b|\bditekankan\b|\bdianjurkan\b",

        # deliverable
        r"\bserahkan\b|\bdiserahkan\b|\bkirimkan\b|\bdikirimkan\b|\bunggah berkas\b|\bupload berkas\b|\bsubmit\b|\bdisubmit\b",
        r"\bmenyertakan lampiran\b|\blampiran lengkap\b|\bkelengkapan berkas\b",

        # jadwal/detail eksekusi
        r"\bmenyusun jadwal\b|\btime line\b|\btimeline\b|\brencana kerja\b|\bRencana Tindak Lanjut\b|\bRTL\b"
    ],

    # ================= LAINNYA =================
    "isu": [
        r"\bisu\b|\bmasalah\b|\bkendala\b|\brisiko\b|\bblokir\b|\bproblem\b|\bcatatan risiko\b"
    ],
    "arahan": [
        r"\barahan\b|\binstruksi\b|\bdiarahkan\b|\bdiminta\b|\bgaris besar\b|\bperhatian\b"
    ],
    "catatan": [
        r"\bcatatan\b|\binformasi\b|\bupdate\b|\bkonteks\b|\bpengantar\b"
    ]
}

def _compile_keysets(custom: Dict[str, List[str]]|None) -> Dict[str, List[re.Pattern]]:
    """Gabungkan base + custom ke regex terkompilasi."""
    merged: Dict[str, List[str]] = {k: list(v) for k, v in BASE_KEYSETS.items()}
    if custom:
        for k, arr in custom.items():
            if not arr: continue
            merged.setdefault(k, [])
            merged[k].extend(arr)
    return {k: [re.compile(p, re.I) for p in v] for k, v in merged.items()}

def _split_candidates(text: str) -> list[str]:
    parts = re.split(r"(?:\n{2,}|(?<=[.!?])\s+)", text)
    return [p.strip("•- \t\r") for p in parts if p and p.strip()]

def _infer_owner(line: str) -> str | None:
    m = RE_OWNER.search(line)
    if m: return m.group(2).strip()
    m2 = re.search(r"\b([A-Z][a-zA-Z_. ]{1,40})\b[^.]{0,30}\b(akan|agar|diminta|ditugaskan)\b", line)
    return m2.group(1).strip() if m2 else None

def _infer_due(line: str) -> str | None:
    m = RE_DATE.search(line)
    return m.group(0) if m else None

def _classify_line(line: str, COMPILED: Dict[str, List[re.Pattern]]) -> tuple[str, float, dict]:
    score = {k: 0.0 for k in COMPILED}
    for k, regs in COMPILED.items():
        for rgx in regs:
            if rgx.search(line): score[k] += 1.0
    if RE_DECISION_VERB.search(line): score["keputusan"] += 1.6
    if RE_ACTION_VERB.search(line):   score["tindak_lanjut"] += 1.1
    if _infer_due(line):              score["tindak_lanjut"] += 0.8
    if _infer_owner(line):            score["tindak_lanjut"] += 0.8

    label = max(score, key=score.get)
    return label, score[label], {
        "text": line.strip(),
        "owner": _infer_owner(line),
        "due_date": _infer_due(line),
        "confidence_local": round(score[label], 2)
    }

def extract_minutes_rule_based(
    transcript: str,
    summary: str | None,
    *,
    max_each: int = 30,
    custom_keywords: Dict[str, List[str]]|None = None
) -> dict:
    """
    Ekstraksi kategori notulen berbasis aturan + kata kunci custom.
    """
    COMPILED = _compile_keysets(custom_keywords)

    if summary and summary.strip() and not summary.strip().startswith("[Gagal"):
        source_text = summary
    else:
        source_text = local_summarize_bullets(transcript, max_sentences=30)

    buckets = {k: [] for k in ["keputusan","tindak_lanjut","isu","arahan","catatan"]}
    ranked: list[tuple[str,float,dict]] = []
    for ln in _split_candidates(source_text):
        label, sc, payload = _classify_line(ln, COMPILED)
        ranked.append((label, sc, payload))

    for label, _, pay in sorted(ranked, key=lambda x: x[1], reverse=True):
        if len(buckets[label]) < max_each and pay["text"] not in [p["text"] for p in buckets[label]]:
            buckets[label].append(pay)

    return buckets

# ===== Notulen (Minutes) Builder =====
def build_minutes_local(transcript: str, summary: str | None, program: str, created_at: str, *, meta: dict|None=None) -> dict:
    """Builder utama untuk halaman HTML lama (section Keputusan/Tindak Lanjut)."""
    try:
        tanggal = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S").strftime("%A, %d %B %Y")
    except Exception:
        tanggal = created_at

    custom = None
    if meta and isinstance(meta, dict):
        custom = {
            "keputusan": meta.get("kw_keputusan", []),
            "tindak_lanjut": meta.get("kw_tindak_lanjut", []),
            "isu": meta.get("kw_isu", []),
            "arahan": meta.get("kw_arahan", []),
            "catatan": meta.get("kw_catatan", []),
        }
        custom = {k:v for k,v in custom.items() if v}

    buckets = extract_minutes_rule_based(transcript, summary, max_each=40, custom_keywords=custom)

    return {
        "title": f"NOTULEN RAPAT {program.upper()}",
        "tanggal": tanggal,
        "agenda": ["Pembukaan", "Pemaparan/Pembahasan", "Keputusan", "Tindak Lanjut", "Penutup"],
        "arahan":        [x["text"] for x in buckets["arahan"]],
        "keputusan":     [x["text"] for x in buckets["keputusan"]],
        "tindak_lanjut": [x["text"] for x in buckets["tindak_lanjut"]],
        "isu":           [x["text"] for x in buckets["isu"]],
        "catatan":       [x["text"] for x in buckets["catatan"]],
    }

def build_minutes_official(transcript: str, summary: str | None, program: str, created_at: str, *, meta: dict|None=None) -> dict:
    """
    Builder untuk format resmi seperti contoh dokumen DPRD:
    I ... IX, dengan fokus pada 'Hasil Rapat' (Keputusan/Tindak Lanjut).
    """
    m = meta or {}
    try:
        hari = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S").strftime("%A")
        tanggal = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S").strftime("%d %B %Y")
    except Exception:
        hari = ""
        tanggal = created_at

    custom = {
        "keputusan": m.get("kw_keputusan", []),
        "tindak_lanjut": m.get("kw_tindak_lanjut", []),
        "isu": m.get("kw_isu", []),
        "arahan": m.get("kw_arahan", []),
        "catatan": m.get("kw_catatan", []),
    }
    custom = {k:v for k,v in custom.items() if v}
    buckets = extract_minutes_rule_based(transcript, summary, max_each=50, custom_keywords=custom)

    return {
        "header": {
            "instansi": (m.get("instansi") or DEFAULT_META["instansi"]).upper(),
            "alamat": m.get("alamat") or DEFAULT_META["alamat"],
            "logo": m.get("logo") or "",
        },
        "laporan": {
            "judul": f"Laporan {program}",
            "jenis": m.get("jenis","Terbuka"),
            "hari": m.get("hari") or hari,
            "tanggal": m.get("tanggal") or tanggal,
            "waktu": m.get("waktu",""),
            "acara": m.get("acara") or f"Rapat {program}",
            "pimpinan": m.get("pimpinan",""),
            "peserta": m.get("peserta", []),
        },
        "hasil": {
            "keputusan": [x["text"] for x in buckets["keputusan"]],
            "tindak_lanjut": [x["text"] for x in buckets["tindak_lanjut"]],
            "isu": [x["text"] for x in buckets["isu"]],
            "arahan": [x["text"] for x in buckets["arahan"]],
            "catatan": [x["text"] for x in buckets["catatan"]],
        },
        "penutup": m.get("penutup","Rapat ditutup pada waktu yang telah ditentukan."),
        "title": f"NOTULEN RAPAT {program.upper()}",
        "tanggal_display": m.get("tanggal") or tanggal
    }

def build_minutes_gpt(transcript: str, summary: str | None, program: str, created_at: str, *, meta: dict|None=None) -> dict:
    """Alias ke builder lokal supaya route lama tetap kompatibel."""
    return build_minutes_local(transcript, summary, program, created_at, meta=meta)

# ==== DOCX builder (format resmi) ====
def build_docx_from_minutes(minutes: dict, meta: dict, tr: dict) -> BytesIO:
    """
    Versi lama (heading Keputusan/Tindak Lanjut) tetap jalan.
    Namun di bawah ini kita bikin DOCX yang tampak seperti template resmi:
    I. Jenis, II. Hari, ... VIII. Hasil Rapat, IX. Penutup
    """
    if not HAVE_DOCX:
        raise RuntimeError("python-docx belum terpasang. Jalankan: pip install python-docx")
    if not isinstance(meta, dict): meta = dict(meta)
    if not isinstance(tr, dict): tr = dict(tr)

    official = build_minutes_official(tr.get("transcript",""), tr.get("summary",""), tr.get("program",""), tr.get("created_at",""), meta=meta)

    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if meta.get("logo"):
        try:
            logo_val = meta["logo"].strip()
            if logo_val.startswith("/uploads/"):
                logo_fs = UPLOAD_DIR / logo_val.split("/uploads/")[-1]
                if logo_fs.exists():
                    doc.add_picture(str(logo_fs), width=Inches(1.0))
                else:
                    log.warning(f"Logo file not found: {logo_fs}")
            else:
                parsed = urlparse(logo_val)
                if parsed.scheme in ("http", "https"):
                    r = requests.get(logo_val, timeout=10)
                    r.raise_for_status()
                    tmp = BytesIO(r.content)
                    doc.add_picture(tmp, width=Inches(1.0))
                else:
                    log.warning(f"Unsupported logo path/URL scheme: {logo_val}")
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            log.warning(f"Gagal memasang logo ke DOCX: {e}")

    kop = (official["header"]["instansi"]).upper()
    alamat = official["header"]["alamat"]

    p = doc.add_paragraph(kop); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: run.bold = True
    p = doc.add_paragraph(alamat); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("".ljust(64, "_")); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run(official["laporan"]["judul"]); r.bold = True; r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_row(k, v):
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    table = doc.add_table(rows=0, cols=2)
    add_row("I. Jenis/Sifat Rapat", official["laporan"]["jenis"] or "—")
    add_row("II. Hari", official["laporan"]["hari"] or "—")
    add_row("III. Tanggal", official["laporan"]["tanggal"] or "—")
    add_row("IV. Waktu", official["laporan"]["waktu"] or "—")
    add_row("V. Acara", official["laporan"]["acara"] or "—")
    add_row("VI. Pimpinan Rapat", official["laporan"]["pimpinan"] or "—")

    doc.add_paragraph()
    doc.add_paragraph("VII. Peserta Rapat")
    peserta = official["laporan"].get("peserta") or []
    if peserta:
        for i, pnama in enumerate(peserta, 1):
            doc.add_paragraph(f"{i}. {pnama}", style=None)
    else:
        doc.add_paragraph("—")

    doc.add_paragraph()
    doc.add_paragraph("VIII. Hasil Rapat")
    hasil = official["hasil"]
    def add_ul(title, items: List[str]):
        if not items: return
        doc.add_paragraph(title).runs[0].bold = True
        for i, it in enumerate(items, 1):
            para = doc.add_paragraph(f"{i}. {it}")
            para.style = doc.styles['List Bullet']

    add_ul("Keputusan:", hasil.get("keputusan", []))
    add_ul("Tindak Lanjut:", hasil.get("tindak_lanjut", []))
    add_ul("Isu/Kendala:", hasil.get("isu", []))
    add_ul("Arahan:", hasil.get("arahan", []))
    add_ul("Catatan:", hasil.get("catatan", []))

    doc.add_paragraph()
    doc.add_paragraph("IX. Penutup")
    doc.add_paragraph(official.get("penutup",""))

    # --- BAGIAN INI YANG DIUBAH UNTUK TANDA TANGAN ---
    for _ in range(3):
        doc.add_paragraph()

    jabatan_ttd = meta.get("ttd_jabatan", "")
    nama_ttd = meta.get("ttd_nama", "")
    pangkat_ttd = meta.get("ttd_pangkat", "")
    nip_ttd = meta.get("ttd_nip", "")

    outer_right_indent = Cm(0.5)
    inner_right_indent = Cm(1.5)

    if jabatan_ttd:
        p_jabatan = doc.add_paragraph(jabatan_ttd)
        p_jabatan.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_jabatan.paragraph_format.right_indent = outer_right_indent
        p_jabatan.paragraph_format.space_before = Pt(0)
        p_jabatan.paragraph_format.space_after = Pt(0)
    else:
        p_jabatan_empty = doc.add_paragraph()
        p_jabatan_empty.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_jabatan_empty.paragraph_format.right_indent = outer_right_indent
        p_jabatan_empty.paragraph_format.space_before = Pt(0)
        p_jabatan_empty.paragraph_format.space_after = Pt(0)

    p_extra_space_after_jabatan = doc.add_paragraph()
    p_extra_space_after_jabatan.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_extra_space_after_jabatan.paragraph_format.right_indent = outer_right_indent
    p_extra_space_after_jabatan.paragraph_format.space_before = Pt(0)
    p_extra_space_after_jabatan.paragraph_format.space_after = Pt(0)

    p_space_before_name = doc.add_paragraph()
    p_space_before_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_space_before_name.paragraph_format.right_indent = inner_right_indent
    p_space_before_name.paragraph_format.space_before = Pt(0)
    p_space_before_name.paragraph_format.space_after = Pt(0)

    if nama_ttd:
        p_nama = doc.add_paragraph()
        run_nama = p_nama.add_run(nama_ttd)
        run_nama.bold = True
        run_nama.underline = True
        p_nama.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_nama.paragraph_format.right_indent = inner_right_indent
        p_nama.paragraph_format.space_before = Pt(0)
        p_nama.paragraph_format.space_after = Pt(2)
    else:
        p_nama_empty = doc.add_paragraph()
        run_empty_name = p_nama_empty.add_run("                    ")
        run_empty_name.underline = True
        p_nama_empty.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_nama_empty.paragraph_format.right_indent = inner_right_indent
        p_nama_empty.paragraph_format.space_before = Pt(0)
        p_nama_empty.paragraph_format.space_after = Pt(2)

    if pangkat_ttd:
        p_pangkat = doc.add_paragraph(pangkat_ttd)
        p_pangkat.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_pangkat.paragraph_format.right_indent = inner_right_indent
        p_pangkat.paragraph_format.space_before = Pt(0)
        p_pangkat.paragraph_format.space_after = Pt(0)
    else:
        p_pangkat_empty = doc.add_paragraph()
        p_pangkat_empty.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_pangkat_empty.paragraph_format.right_indent = inner_right_indent
        p_pangkat_empty.paragraph_format.space_before = Pt(0)
        p_pangkat_empty.paragraph_format.space_after = Pt(0)

    if nip_ttd:
        p_nip = doc.add_paragraph()
        run_nip = p_nip.add_run(f"NIP. {nip_ttd}")
        run_nip.bold = True
        p_nip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_nip.paragraph_format.right_indent = outer_right_indent
        p_nip.paragraph_format.space_before = Pt(0)
        p_nip.paragraph_format.space_after = Pt(0)
    else:
        p_nip_empty = doc.add_paragraph()
        p_nip_empty.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_nip_empty.paragraph_format.right_indent = outer_right_indent
        p_nip_empty.paragraph_format.space_before = Pt(0)
        p_nip_empty.paragraph_format.space_after = Pt(0)
    # --- AKHIR PERUBAHAN TANDA TANGAN ---

    bio = BytesIO(); doc.save(bio); bio.seek(0)
    return bio

def get_now_playing() -> str:
    try:
        resp = requests.get("https://admin.sebayu.my.id/api/nowplaying/sebayu", timeout=6)
        data = resp.json()
        song = data["now_playing"]["song"]
        artist = song.get("artist", ""); title = song.get("title", "")
        live = data["now_playing"].get("live"); listeners = data.get("listeners", {}).get("current", 0)
        if live and live.get("is_live", False):
            return f"🎙️ Sedang live oleh {live.get('streamer_name','penyiar')} – memutar {title} {artist}".strip()
        return f"🎶 Sekarang memutar: {title} {artist} | 👥 {listeners} pendengar".strip()
    except Exception as e:
        return f"[Gagal ambil status: {e}]"

def set_progress(job_id: str, pct: int, msg: str, *, done: bool=False, error: str|None=None, tid: int|None=None):
    PROGRESS[job_id] = {"pct": int(pct), "msg": msg, "done": done, "error": error, "tid": tid}
    log.info(f"[{job_id}] {pct}% {msg}")

def run_transcribe_job(job_id: str, save_path: Path, program: str, mode: str, manual_choice: str, do_chunk: bool, do_summary: bool):
    try:
        set_progress(job_id, 10, "Mulai proses")
        full_text = transcribe_audio_pipeline(
            save_path, mode=mode, manual_choice=manual_choice, do_chunk=do_chunk,
            progress=lambda p,m: set_progress(job_id, p, m)
        )
        summary_text = None
        if do_summary:
            set_progress(job_id, 92, "Merangkum (lokal)…")
            try:
                summary_text = local_summarize_bullets(full_text)
            except Exception as e:
                summary_text = f"[Gagal merangkum: {e}]"

        set_progress(job_id, 98, "Menyimpan ke database")
        with get_db() as db:
            cur = db.execute(
                "INSERT INTO transcripts(program, filename, transcript, created_at, summary) VALUES(?,?,?,?,?)",
                (program, save_path.name, full_text, now_str(), summary_text),
            )
            tid = cur.lastrowid

            # --- Bersihkan teks & simpan ke kolom cleaned_transcript (jika ada) ---
            try:
                cleaned = clean_text_id(full_text)
                db.execute("UPDATE transcripts SET cleaned_transcript=? WHERE id=?", (cleaned, tid))
            except Exception as e:
                log.warning(f"Gagal simpan cleaned_transcript (abaikan jika kolom belum ada): {e}")
            # --- END ---

            db.commit()

        set_progress(job_id, 100, "Selesai ✅", done=True, tid=tid)
    except Exception as e:
        log.exception("Transcribe job error")
        set_progress(job_id, 100, f"Gagal: {e}", done=True, error=str(e))

def handle_chat_message(text: str) -> str:
    t = (text or "").strip()
    if not t:
        return "Halo! Ketik 'jadwal', 'siaran', 'lagu sekarang', atau 'request Judul - Artis'."
    lower = t.lower()
    if lower in {"help", "bantuan", "/start"}:
        return (
            "Perintah yang tersedia:\n"
            "• jadwal — Lihat jadwal siaran hari ini\n"
            "• siaran — Info program terjadwal sekarang\n"
            "• lagu sekarang / status — Info lagu real-time dari Azuracast\n"
            "• request Judul - Artis — Kirim request lagu\n"
            "• help — Bantuan"
        )
    if "jadwal" in lower: return get_today_schedule_text()
    if "siaran" in lower:
        cp = current_program()
        if cp:
            prog, host = cp; host_s = f" oleh {host}" if host else ""
            return f"📻 Program terjadwal sekarang: {prog}{host_s}"
        return "Tidak ada program terjadwal saat ini."
    if "lagu" in lower or "status" in lower: return get_now_playing()
    if lower.startswith("request ") or lower.startswith("req "):
        return "Oke, aku catat. (Akan tersimpan saat kamu kirim.)"
    return "Maaf, aku belum paham. Ketik 'help' untuk bantuan."
